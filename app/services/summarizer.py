
from __future__ import annotations

from io import BytesIO
from typing import List, Tuple, Optional

from docx import Document
from openai import OpenAI

from app.services.pdf_utils import extract_text_from_pdf_bytes
from app.services.docx_writer import write_text_to_docx_bytes
from app.core.config import settings

import logging

# Optional token counting
try:
    import tiktoken
except Exception:
    tiktoken = None

client = OpenAI(api_key=settings.OPENAI_API_KEY)

def _count_tokens(text: str, model: str) -> int:
    if not tiktoken:
        return max(1, len(text) // 4)  # crude fallback
    try:
        enc = tiktoken.encoding_for_model(model)
    except Exception:
        enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))

def _read_docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text)

def extract_template_instructions(docx_stream: BytesIO) -> str:
    doc = Document(docx_stream)
    return "\n".join(p.text for p in doc.paragraphs)

def _chunk_text(text: str, size: int, overlap: int) -> List[str]:
    if len(text) <= size:
        return [text]
    out, i = [], 0
    step = max(1, size - overlap)
    while i < len(text):
        out.append(text[i:i+size])
        i += step
    return out

def _chat_once(prompt: str, model: str) -> str:
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
    )
    return resp.choices[0].message.content.strip()

def _summarize_chunks(text: str, instructions: Optional[str]) -> Tuple[str, int, int]:
    model = settings.OPENAI_MODEL
    chunks = _chunk_text(text, size=settings.CHUNK_SIZE, overlap=settings.CHUNK_OVERLAP)
    tin = tout = 0

    # map
    partials = []
    for ch in chunks:
        prompt = f"Please summarize the following text.\n\n{ch}\n\nSummary:"
        tin += _count_tokens(prompt, model)
        s = _chat_once(prompt, model)
        tout += _count_tokens(s, model)
        partials.append(s)

    # reduce
    combined = "\n\n".join(partials)
    if instructions:
        reduce_prompt = (
            "Using the following instructions, create a concise, structured summary of the material.\n\n"
            f"Instructions:\n{instructions}\n\n"
            f"Material:\n{combined}\n\n"
            "Final summary:"
        )
    else:
        reduce_prompt = f"Create a concise, structured summary of the following material.\n\n{combined}\n\nFinal summary:"
    tin += _count_tokens(reduce_prompt, model)
    final = _chat_once(reduce_prompt, model)
    tout += _count_tokens(final, model)
    return final, tin, tout

def _combine_across_files(file_summaries: List[Tuple[str, str]], instructions: Optional[str]) -> Tuple[str, int, int]:
    model = settings.OPENAI_MODEL
    combined_text = "".join([f"Summary of {n}:\n{s}\n\n" for n, s in file_summaries])
    if instructions:
        prompt = (
            "Given the template instructions and summarized text, arrange the content per the template order.\n"
            "Do not significantly rephrase. Bold headings using **like this**. If the template specifies sub-sections,"
            " list them using a), b), c).\n\n"
            f"Template instructions:\n{instructions}\n\nSummarized text:\n{combined_text}\n\nFinal arranged document:"
        )
    else:
        prompt = (
            "Combine the following summaries into one coherent document with clear section headings. "
            "Bold headings using **like this**.\n\n"
            f"{combined_text}\n\nFinal document:"
        )
    tin = _count_tokens(prompt, model)
    final = _chat_once(prompt, model)
    tout = _count_tokens(final, model)
    return final, tin, tout

def _read_any_text(filename: str, data: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf_bytes(data)
    if ext == "docx":
        return _read_docx_text(data)
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""

def summarize_many_documents_into_one(
    files: List[Tuple[str, bytes]],
    instructions: Optional[str] = None,
) -> Tuple[bytes, dict]:
    if not settings.OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY not set")

    per_file: List[Tuple[str, str]] = []
    total_in = total_out = 0

    for fname, data in files:
        raw = _read_any_text(fname, data)
        if not raw.strip():
            logging.warning(f"{fname}: empty or unreadable content; skipping.")
            continue
        summ, tin, tout = _summarize_chunks(raw, instructions=None)
        per_file.append((fname, summ))
        total_in += tin; total_out += tout

    if not per_file:
        raise RuntimeError("No readable inputs.")

    final_text, cin, cout = _combine_across_files(per_file, instructions)
    total_in += cin; total_out += cout

    docx_bytes = write_text_to_docx_bytes(final_text)
    return docx_bytes, {"input_tokens": total_in, "output_tokens": total_out, "total_tokens": total_in + total_out}

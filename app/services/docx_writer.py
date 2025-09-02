
from io import BytesIO
from docx import Document
from docx.shared import Pt

def _write_markdownish_text(doc: Document, text: str):
    """
    Very light '**bold**' support: split on ** and toggle bold.
    """
    for block in text.split("\n\n"):
        p = doc.add_paragraph()
        bold = False
        for segment in block.split("**"):
            r = p.add_run(segment)
            r.font.size = Pt(11)
            r.bold = bold
            bold = not bold
    return doc

def write_text_to_docx_bytes(full_text: str) -> bytes:
    doc = Document()
    _write_markdownish_text(doc, full_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
